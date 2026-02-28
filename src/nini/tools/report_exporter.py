"""报告导出器 — 支持 Markdown 转 DOCX 和 PDF。

提供发表级报告的多格式导出能力。
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class ReportExporter:
    """报告导出器基类。"""

    def __init__(self, markdown_content: str) -> None:
        self.markdown = markdown_content

    def export(self) -> bytes:
        """导出为字节数据，子类实现。"""
        raise NotImplementedError


class DOCXExporter(ReportExporter):
    """Markdown 转 DOCX 导出器。"""

    def __init__(self, markdown_content: str, title: str = "科研数据分析报告") -> None:
        super().__init__(markdown_content)
        self.title = title

    def export(self) -> bytes:
        """导出为 DOCX 格式。
        
        Returns:
            DOCX 文件字节数据
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx 未安装，无法导出 DOCX")
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # 解析 Markdown 并转换为 DOCX
        lines = self.markdown.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # 标题
            if stripped.startswith('# '):
                # 一级标题 - 报告标题
                heading = doc.add_heading(stripped[2:], level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif stripped.startswith('## '):
                # 二级标题
                doc.add_heading(stripped[3:], level=1)
                
            elif stripped.startswith('### '):
                # 三级标题
                doc.add_heading(stripped[4:], level=2)
                
            elif stripped.startswith('#### '):
                # 四级标题
                doc.add_heading(stripped[5:], level=3)
                
            # 引用块（元信息）
            elif stripped.startswith('>'):
                p = doc.add_paragraph()
                run = p.add_run(stripped[1:].strip())
                run.italic = True
                run.font.color.rgb = None  # 灰色由样式控制
                
            # 列表
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
                
            elif re.match(r'^\d+\. ', stripped):
                # 有序列表
                text = re.sub(r'^\d+\. ', '', stripped)
                doc.add_paragraph(text, style='List Number')
                
            # 表格（简化处理）
            elif stripped.startswith('|'):
                # 跳过分隔行
                if not stripped.replace('|', '').replace('-', '').replace(':', '').strip():
                    i += 1
                    continue
                    
                # 解析表格行
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if cells:
                    # 简单表格处理：作为带缩进的段落
                    p = doc.add_paragraph()
                    p.add_run(' | '.join(cells)).bold = True
                    
            # 代码块
            elif stripped.startswith('```'):
                # 收集代码块内容
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    
            # 空行
            elif not stripped:
                doc.add_paragraph()
                
            # 普通段落
            else:
                # 处理行内格式
                p = doc.add_paragraph()
                self._add_formatted_text(p, stripped)
                
            i += 1
        
        # 保存到字节流
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _add_formatted_text(self, paragraph: Any, text: str) -> None:
        """添加带格式的文本到段落。
        
        处理 **粗体**、*斜体*、`代码` 等行内格式。
        """
        import re
        
        # 分割格式标记
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗体
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                # 斜体
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # 代码
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # 普通文本
                paragraph.add_run(part)


class PDFExporter(ReportExporter):
    """Markdown 转 PDF 导出器。"""

    def __init__(
        self,
        markdown_content: str,
        title: str = "科研数据分析报告",
        journal_style: str = "default",
    ) -> None:
        super().__init__(markdown_content)
        self.title = title
        self.journal_style = journal_style

    def export(self) -> bytes:
        """导出为 PDF 格式。
        
        Returns:
            PDF 文件字节数据
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            logger.error("reportlab 未安装，无法导出 PDF")
            raise ImportError("请安装 reportlab: pip install reportlab")

        # 尝试注册中文字体
        try:
            # 常见中文字体路径
            font_paths = [
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/System/Library/Fonts/PingFang.ttc",  # macOS
                "C:/Windows/Fonts/simhei.ttf",  # Windows
                "C:/Windows/Fonts/simsun.ttc",
            ]
            
            chinese_font = None
            for fp in font_paths:
                if Path(fp).exists():
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', fp))
                        chinese_font = 'Chinese'
                        break
                    except Exception:
                        continue
            
            if chinese_font is None:
                logger.warning("未找到中文字体，PDF 中文显示可能异常")
                chinese_font = 'Helvetica'
        except Exception:
            chinese_font = 'Helvetica'

        # 创建 PDF 文档
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # 定义样式
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=18,
            alignment=1,  # 居中
            spaceAfter=30,
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=14,
            spaceAfter=12,
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=12,
            spaceAfter=10,
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
        
        code_style = ParagraphStyle(
            'CustomCode',
            parent=styles['Code'],
            fontName='Courier',
            fontSize=9,
            leftIndent=20,
            textColor=colors.darkblue,
        )

        # 解析 Markdown
        story = []
        lines = self.markdown.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # 标题
            if stripped.startswith('# '):
                story.append(Paragraph(stripped[2:], title_style))
                story.append(Spacer(1, 0.2 * inch))
                
            elif stripped.startswith('## '):
                story.append(Paragraph(stripped[3:], heading1_style))
                story.append(Spacer(1, 0.1 * inch))
                
            elif stripped.startswith('### '):
                story.append(Paragraph(stripped[4:], heading2_style))
                
            # 引用块
            elif stripped.startswith('>'):
                p = Paragraph(f"<i>{self._escape_xml(stripped[1:].strip())}</i>", normal_style)
                story.append(p)
                
            # 列表
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = '• ' + self._escape_xml(stripped[2:])
                story.append(Paragraph(text, normal_style))
                
            # 代码块
            elif stripped.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    story.append(Paragraph(self._escape_xml(code_text), code_style))
                    
            # 空行
            elif not stripped:
                story.append(Spacer(1, 0.1 * inch))
                
            # 普通段落
            else:
                # 转换 Markdown 格式为 HTML
                html_text = self._markdown_to_html(stripped)
                story.append(Paragraph(html_text, normal_style))
                
            i += 1

        # 生成 PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _escape_xml(self, text: str) -> str:
        """转义 XML 特殊字符。"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))
    
    def _markdown_to_html(self, text: str) -> str:
        """将 Markdown 行内格式转为 HTML。
        
        处理 **粗体**、*斜体*、`代码`。
        """
        import re
        
        text = self._escape_xml(text)
        
        # 粗体 **text** -> <b>text</b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        
        # 斜体 *text* -> <i>text</i>
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        
        # 代码 `text` -> <font face="Courier">text</font>
        text = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', text)
        
        return text


def export_report(
    markdown_content: str,
    format: str,  # noqa: A002
    title: str = "科研数据分析报告",
    journal_style: str = "default",
) -> bytes:
    """导出报告为指定格式。
    
    Args:
        markdown_content: Markdown 格式的报告内容
        format: 导出格式（docx/pdf）
        title: 报告标题
        journal_style: 期刊风格
        
    Returns:
        导出文件的字节数据
        
    Raises:
        ValueError: 不支持的格式
        ImportError: 缺少必要的依赖
    """
    format = format.lower()
    
    if format == "docx":
        exporter = DOCXExporter(markdown_content, title)
        return exporter.export()
    
    elif format == "pdf":
        exporter = PDFExporter(markdown_content, title, journal_style)
        return exporter.export()
    
    else:
        raise ValueError(f"不支持的导出格式: {format}")
